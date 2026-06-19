"""文本文件格式转换器 —— 将各种文本文件转为 .md，并校验转换质量"""

import re
from pathlib import Path
from dataclasses import dataclass

# 支持转换的格式
SUPPORTED_EXTENSIONS = {
    ".docx", ".doc",   # Word
    ".html", ".htm",   # 网页
    ".rtf",            # 富文本
    ".odt",            # OpenDocument
    ".tex",            # LaTeX
    ".csv",            # 表格
    ".json",           # JSON
    ".xml",            # XML
    ".epub",           # 电子书
    ".mobi",           # Kindle 电子书
}

# 已经可以直接用的格式（无需转换）
SKIP_EXTENSIONS = {".txt", ".md", ".text", ".pdf"}


@dataclass
class ConvertResult:
    """转换结果"""
    success: bool
    source_path: str
    output_path: str | None
    message: str
    quality_score: float = 0.0  # 0-1，转换质量评分
    issues: list[str] | None = None


def convert_to_md(source_path: str, output_dir: str | None = None) -> ConvertResult:
    """将文本文件转换为 .md 格式

    Args:
        source_path: 源文件路径
        output_dir: 输出目录，默认与源文件同目录

    Returns:
        ConvertResult 转换结果
    """
    src = Path(source_path)

    if not src.exists():
        return ConvertResult(False, source_path, None, f"文件不存在: {source_path}")

    ext = src.suffix.lower()

    if ext in SKIP_EXTENSIONS:
        return ConvertResult(False, source_path, None, f"文件格式 {ext} 无需转换")

    if ext not in SUPPORTED_EXTENSIONS:
        return ConvertResult(False, source_path, None,
                             f"不支持的格式: {ext}，支持: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")

    # 确定输出路径
    out_dir = Path(output_dir) if output_dir else src.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{src.stem}.md"

    # 根据格式分派转换器
    converters = {
        ".docx": _convert_docx,
        ".doc": _convert_doc,
        ".html": _convert_html,
        ".htm": _convert_html,
        ".rtf": _convert_rtf,
        ".odt": _convert_odt,
        ".tex": _convert_tex,
        ".csv": _convert_csv,
        ".json": _convert_json,
        ".xml": _convert_xml,
        ".epub": _convert_epub,
        ".mobi": _convert_mobi,
    }

    converter = converters.get(ext)
    if not converter:
        return ConvertResult(False, source_path, None, f"未找到 {ext} 的转换器")

    try:
        md_content = converter(src)
    except Exception as e:
        return ConvertResult(False, source_path, None, f"转换失败: {e}")

    if not md_content or not md_content.strip():
        return ConvertResult(False, source_path, None, "转换结果为空")

    # 校验转换质量
    quality_score, issues = _validate_quality(md_content)

    # 写入文件
    try:
        out_path.write_text(md_content, encoding="utf-8")
    except Exception as e:
        return ConvertResult(False, source_path, None, f"写入文件失败: {e}")

    msg = f"转换成功: {src.name} → {out_path.name}"
    if issues:
        msg += f"（质量评分: {quality_score:.0%}，发现 {len(issues)} 个潜在问题）"

    return ConvertResult(
        success=True,
        source_path=source_path,
        output_path=str(out_path),
        message=msg,
        quality_score=quality_score,
        issues=issues if issues else None,
    )


def batch_convert(input_dir: str, output_dir: str | None = None) -> list[ConvertResult]:
    """批量转换目录下所有支持格式的文件

    Args:
        input_dir: 输入目录
        output_dir: 输出目录，默认为输入目录下的 converted/ 子目录

    Returns:
        转换结果列表
    """
    input_path = Path(input_dir)
    if not input_path.exists():
        return [ConvertResult(False, input_dir, None, f"目录不存在: {input_dir}")]

    if output_dir is None:
        output_dir = str(input_path / "converted")

    results = []
    for filepath in sorted(input_path.rglob("*")):
        if filepath.is_file() and filepath.suffix.lower() in SUPPORTED_EXTENSIONS:
            result = convert_to_md(str(filepath), output_dir)
            results.append(result)
            status = "✓" if result.success else "✗"
            print(f"  {status} {result.message}")

    return results


# ==================== 各格式转换器 ====================


def _convert_docx(src: Path) -> str:
    """Word (.docx) → Markdown"""
    try:
        import docx
    except ImportError:
        raise RuntimeError("需要安装 python-docx: pip install python-docx")

    doc = docx.Document(str(src))
    md_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_parts.append("")
            continue

        style = para.style.name.lower() if para.style else ""

        if "heading 1" in style or "标题 1" in style:
            md_parts.append(f"# {text}")
        elif "heading 2" in style or "标题 2" in style:
            md_parts.append(f"## {text}")
        elif "heading 3" in style or "标题 3" in style:
            md_parts.append(f"### {text}")
        elif "heading" in style or "标题" in style:
            md_parts.append(f"#### {text}")
        elif "list" in style or "bullet" in style:
            md_parts.append(f"- {text}")
        else:
            # 处理加粗/斜体
            formatted = _extract_runs_formatting(para)
            md_parts.append(formatted)

    # 处理表格
    for table in doc.tables:
        md_parts.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            md_parts.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
        md_parts.append("")

    return "\n".join(md_parts)


def _extract_runs_formatting(para) -> str:
    """提取段落中的加粗/斜体格式"""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts) if parts else para.text


def _convert_doc(src: Path) -> str:
    """旧版 Word (.doc) → Markdown"""
    # .doc 格式需要额外工具，尝试用 antiword 或 fallback
    try:
        import subprocess
        result = subprocess.run(["antiword", str(src)], capture_output=True, text=True, timeout=30)
        if result.returncode == 0:
            return _plain_text_to_md(result.stdout)
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: 提示用户转换为 docx
    raise RuntimeError(
        f"无法转换 .doc 格式，请先用 Word 另存为 .docx 格式: {src.name}"
    )


def _convert_html(src: Path) -> str:
    """HTML → Markdown"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("需要安装 beautifulsoup4: pip install beautifulsoup4")

    # 尝试不同编码
    content = _read_with_fallback(src)
    soup = BeautifulSoup(content, "html.parser")

    # 移除 script 和 style
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()

    md_parts = []

    for elem in soup.body.descendants if soup.body else soup.descendants:
        if not hasattr(elem, "name") or not elem.name:
            continue

        text = elem.get_text(strip=True)
        if not text:
            continue

        tag = elem.name.lower()
        if tag in ("h1",):
            md_parts.append(f"\n# {text}\n")
        elif tag in ("h2",):
            md_parts.append(f"\n## {text}\n")
        elif tag in ("h3",):
            md_parts.append(f"\n### {text}\n")
        elif tag in ("h4", "h5", "h6"):
            md_parts.append(f"\n#### {text}\n")
        elif tag == "p":
            md_parts.append(f"\n{text}\n")
        elif tag == "li":
            md_parts.append(f"- {text}")
        elif tag == "blockquote":
            md_parts.append(f"\n> {text}\n")
        elif tag == "pre":
            md_parts.append(f"\n```\n{text}\n```\n")
        elif tag == "td" or tag == "th":
            pass  # 表格在 table 层处理

    # 简单的表格处理
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if rows:
            md_parts.append("")
            for i, row in enumerate(rows):
                cells = [td.get_text(strip=True).replace("|", "\\|")
                         for td in row.find_all(["td", "th"])]
                if cells:
                    md_parts.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        md_parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
            md_parts.append("")

    result = "\n".join(md_parts)
    # 清理多余空行
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


def _convert_rtf(src: Path) -> str:
    """RTF → Markdown"""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise RuntimeError("需要安装 striprtf: pip install striprtf")

    content = _read_with_fallback(src)
    plain_text = rtf_to_text(content)
    return _plain_text_to_md(plain_text)


def _convert_odt(src: Path) -> str:
    """ODT → Markdown"""
    try:
        from odf import text, teletype
        from odf.opendocument import load
    except ImportError:
        raise RuntimeError("需要安装 odfpy: pip install odfpy")

    doc = load(str(src))
    md_parts = []

    for elem in doc.getElementsByType(text.P):
        para_text = teletype.extractText(elem).strip()
        if para_text:
            md_parts.append(para_text)
        else:
            md_parts.append("")

    return "\n\n".join(md_parts)


def _convert_tex(src: Path) -> str:
    """LaTeX → Markdown（基础转换）"""
    content = _read_with_fallback(src)

    # 移除 preamble
    content = re.sub(r'\\documentclass.*?(?=\\begin)', '', content, flags=re.DOTALL)
    content = re.sub(r'\\usepackage.*?\n', '', content)
    content = re.sub(r'\\begin\{document\}', '', content)
    content = re.sub(r'\\end\{document\}', '', content)

    # 标题
    content = re.sub(r'\\section\{(.*?)\}', r'\n# \1\n', content)
    content = re.sub(r'\\subsection\{(.*?)\}', r'\n## \1\n', content)
    content = re.sub(r'\\subsubsection\{(.*?)\}', r'\n### \1\n', content)

    # 格式
    content = re.sub(r'\\textbf\{(.*?)\}', r'**\1**', content)
    content = re.sub(r'\\textit\{(.*?)\}', r'*\1*', content)
    content = re.sub(r'\\emph\{(.*?)\}', r'*\1*', content)

    # 列表
    content = re.sub(r'\\begin\{itemize\}', '', content)
    content = re.sub(r'\\end\{itemize\}', '', content)
    content = re.sub(r'\\item\s*', '- ', content)

    # 清理剩余命令
    content = re.sub(r'\\[a-zA-Z]+\*?\{([^}]*)\}', r'\1', content)
    content = re.sub(r'\\[a-zA-Z]+\*?', '', content)
    content = re.sub(r'[{}]', '', content)

    return content.strip()


def _convert_csv(src: Path) -> str:
    """CSV → Markdown 表格"""
    import csv
    import io

    content = _read_with_fallback(src)
    reader = csv.reader(io.StringIO(content))
    rows = list(reader)

    if not rows:
        return ""

    md_parts = []
    # 表头
    header = rows[0]
    md_parts.append("| " + " | ".join(h.replace("|", "\\|") for h in header) + " |")
    md_parts.append("| " + " | ".join(["---"] * len(header)) + " |")

    # 数据行
    for row in rows[1:]:
        cells = [cell.replace("|", "\\|") for cell in row]
        # 补齐列数
        while len(cells) < len(header):
            cells.append("")
        md_parts.append("| " + " | ".join(cells[:len(header)]) + " |")

    return "\n".join(md_parts)


def _convert_json(src: Path) -> str:
    """JSON → Markdown"""
    import json

    content = _read_with_fallback(src)
    data = json.loads(content)

    if isinstance(data, list):
        # JSON 数组 → 表格
        if data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            md_parts = []
            md_parts.append("| " + " | ".join(str(h) for h in headers) + " |")
            md_parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for item in data:
                cells = [str(item.get(h, "")).replace("|", "\\|") for h in headers]
                md_parts.append("| " + " | ".join(cells) + " |")
            return "\n".join(md_parts)
        else:
            return "\n".join(f"- {item}" for item in data)
    elif isinstance(data, dict):
        # JSON 对象 → 键值对
        md_parts = []
        for k, v in data.items():
            if isinstance(v, (list, dict)):
                md_parts.append(f"## {k}\n\n```json\n{json.dumps(v, ensure_ascii=False, indent=2)}\n```")
            else:
                md_parts.append(f"**{k}:** {v}")
        return "\n\n".join(md_parts)
    else:
        return str(data)


def _convert_xml(src: Path) -> str:
    """XML → Markdown"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("需要安装 beautifulsoup4: pip install beautifulsoup4")

    content = _read_with_fallback(src)
    soup = BeautifulSoup(content, "xml")

    # 递归提取文本结构
    md_parts = []
    for elem in soup.descendants:
        if not hasattr(elem, "name") or not elem.name:
            continue
        text = elem.get_text(strip=True)
        if not text or text == elem.name:
            continue
        # 只处理叶子节点
        if not elem.find():
            md_parts.append(f"**{elem.name}:** {text}")

    return "\n\n".join(md_parts) if md_parts else soup.get_text()


def _is_book_noise(text: str) -> bool:
    """判断 EPUB 单元是否为非正文噪声（目录、版权页、制作说明等）

    保守策略：只过滤明显的非内容项，保留序言/前言/导读。
    """
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return True

    joined = '\n'.join(lines)
    total_chars = len(joined)

    # ── 关键词噪声：前 200 字内出现以下关键词组合 ──
    head = joined[:200].lower()

    # 制作说明 / 版本历史 / 格式说明
    _format_keywords = ['制作说明', '版本历史', '制作信息', '排版说明',
                        '阅读效果', '多看设置', 'sigil', 'calibre',
                        'kindle', 'epub制作', '电子书制作']
    if any(kw in head for kw in _format_keywords):
        return True

    # 出版信息 / 版权页
    _copyright_keywords = ['isbn', 'isbn：', 'isbn:', 'cip数据',
                           '图书在版编目', '版权信息', '版权所有',
                           '出版社：', '出版社:', '出版日期',
                           '印刷:', '印刷：', '开本:', '开本：',
                           '字数:', '字数：', '定价:', '定价：']
    _copyright_hits = sum(1 for kw in _copyright_keywords if kw in joined.lower())
    if _copyright_hits >= 2:
        return True

    # ── 目录检测：大量短行 + 章节标题模式 ──
    if len(lines) >= 8:
        short_lines = [l for l in lines if len(l) < 40]
        short_ratio = len(short_lines) / len(lines)
        # 超过 70% 是短行，且包含"目录"或章节模式
        if short_ratio > 0.7:
            if any(kw in head for kw in ['目录', '总目录', '目  录', '目 录']):
                return True
            # 纯目录模式：每行都是"第X章" / "XX第一" / 数字编号
            toc_pattern = re.compile(
                r'^(第.{1,6}[章卷篇回节部]|'
                r'[一-鿿]{2,6}第[一二三四五六七八九十百千]+|'
                r'\d{1,3}[\.\、]|'
                r'[一二三四五六七八九十百千]+[\.\、])')
            toc_lines = [l for l in lines if toc_pattern.match(l) or len(l) < 25]
            if len(toc_lines) / len(lines) > 0.6:
                return True

    # ── 极短内容（< 80 字）且含噪声关键词 ──
    if total_chars < 80:
        _short_noise = ['目录', '版权', '封面', '扉页', '书名页',
                        '出版', '印刷', '发行', '经销']
        if any(kw in joined for kw in _short_noise):
            return True

    return False


def _convert_epub(src: Path) -> str:
    """EPUB → Markdown"""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        raise RuntimeError("需要安装 ebooklib beautifulsoup4: pip install ebooklib beautifulsoup4")

    book = epub.read_epub(str(src))
    md_parts = []

    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            html_content = item.get_content().decode("utf-8", errors="ignore")
            soup = BeautifulSoup(html_content, "html.parser")
            text = soup.get_text(separator="\n").strip()
            if text and not _is_book_noise(text):
                md_parts.append(text)

    return "\n\n---\n\n".join(md_parts)


def _convert_mobi(src: Path) -> str:
    """MOBI (Kindle) → Markdown（解压后复用 HTML 转换逻辑）"""
    try:
        import mobi
    except ImportError:
        raise RuntimeError("需要安装 mobi: pip install mobi")

    import shutil
    tempdir, filepath = mobi.extract(str(src))

    try:
        # mobi 解压后得到 HTML 文件，复用 _convert_html 逻辑
        html_path = Path(filepath)
        if html_path.suffix.lower() in ('.html', '.htm'):
            return _convert_html(html_path)

        # 如果不是 HTML，查找目录中的 HTML 文件
        html_files = list(Path(tempdir).rglob("*.html")) + list(Path(tempdir).rglob("*.htm"))
        if html_files:
            return _convert_html(html_files[0])

        # 最后 fallback：读取文本内容
        content = html_path.read_text(encoding="utf-8", errors="ignore")
        return _plain_text_to_md(content)
    finally:
        shutil.rmtree(tempdir, ignore_errors=True)


# ==================== 工具函数 ====================


def _read_with_fallback(src: Path) -> str:
    """尝试多种编码读取文件"""
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030", "big5", "latin-1"]
    for enc in encodings:
        try:
            return src.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    # 最终 fallback
    return src.read_bytes().decode("utf-8", errors="replace")


def _plain_text_to_md(text: str) -> str:
    """纯文本转 Markdown（尝试识别标题和段落）"""
    lines = text.split("\n")
    md_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            md_lines.append("")
            continue

        # 短行全大写/全中文大写 → 可能是标题
        if len(stripped) < 50 and stripped == stripped.upper() and len(stripped) > 2:
            md_lines.append(f"## {stripped.title()}")
        else:
            md_lines.append(stripped)

    result = "\n".join(md_lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()


# ==================== 质量校验 ====================


def _validate_quality(content: str) -> tuple[float, list[str]]:
    """校验转换后的 Markdown 内容质量

    Returns:
        (quality_score, issues) - score 范围 0-1，issues 为问题描述列表
    """
    issues = []
    score = 1.0

    # 1. 长度检查
    if len(content) < 50:
        issues.append("内容过短（< 50 字），可能转换不完整")
        score -= 0.4

    # 2. 乱码检测：连续非常用字符
    garbled_pattern = re.compile(r'[^\x00-\x7F一-鿿　-〿＀-￯'
                                 r'\n\r\t .,;:!?()[]{}"""\'\-—…、。，；：！？（）【】《》\n]{5,}')
    garbled_matches = garbled_pattern.findall(content)
    if garbled_matches:
        issues.append(f"疑似乱码片段: {garbled_matches[:3]}")
        score -= 0.3 * min(len(garbled_matches), 3)

    # 3. 乱码比例：非中英文+标点的字符占比
    total_chars = len(content)
    if total_chars > 0:
        normal_chars = len(re.findall(r'[\x00-\x7F一-鿿　-〿＀-￯\n\r\t .,;:!?]', content))
        normal_ratio = normal_chars / total_chars
        if normal_ratio < 0.7:
            issues.append(f"正常字符比例偏低: {normal_ratio:.0%}，可能存在乱码")
            score -= 0.3

    # 4. 可读性：连续无标号的长行（可能是未正确切分的内容）
    long_lines = [i for i, line in enumerate(content.split("\n"))
                  if len(line) > 200 and not line.startswith("#") and not line.startswith("|")]
    if len(long_lines) > 5:
        issues.append(f"有 {len(long_lines)} 行超长文本（>200字），可能未正确切分")
        score -= 0.1

    # 5. 编码残留：HTML 实体
    html_entities = re.findall(r'&[a-zA-Z]+;|&#\d+;', content)
    if html_entities:
        unique_entities = list(set(html_entities))[:5]
        issues.append(f"残留 HTML 实体: {unique_entities}")
        score -= 0.1

    # 6. 空行过多
    if "\n\n\n\n" in content:
        issues.append("存在连续多余空行")
        score -= 0.05

    score = max(0.0, min(1.0, score))
    return score, issues
